import docx

doc = docx.Document(r'c:/Users/issuser/Documents/xwechat_files/wxid_f384e2oubb6p22_6fbf/msg/file/2026-07/广汽集团零采一体化项目-价格目录、工装价格库CE权限_用户测试用例V1(1).docx')

with open(r'c:\Users\issuser\daily-todo\test_output.txt', 'w', encoding='utf-8') as f:
    f.write('=== 段落 ===\n')
    for i, p in enumerate(doc.paragraphs):
        if p.text.strip():
            f.write(f'[{i}] style={p.style.name} text={p.text[:200]}\n')

    for ti, t in enumerate(doc.tables):
        f.write(f'\n=== 表格{ti}: {len(t.rows)}行 x {len(t.columns)}列 ===\n')
        for r, row in enumerate(t.rows):
            cells = [c.text.replace('\n', '\\n') for c in row.cells]
            f.write(f'  Row{r}: | {" | ".join(cells)}\n')
print('done')
