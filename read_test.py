import docx, json

doc = docx.Document(r'c:/Users/issuser/Documents/xwechat_files/wxid_f384e2oubb6p22_6fbf/msg/file/2026-07/广汽集团零采一体化项目-价格目录、工装价格库CE权限_用户测试用例V1(1).docx')

# 先看所有段落
print("=== 段落 ===")
for i, p in enumerate(doc.paragraphs):
    if p.text.strip():
        print(f'[{i}] style={p.style.name} text={p.text[:150]}')

# 看表格
for ti, t in enumerate(doc.tables):
    print(f'\n=== 表格{ti}: {len(t.rows)}行 x {len(t.columns)}列 ===')
    for r, row in enumerate(t.rows):
        cells = [c.text.replace('\n', '\\n') for c in row.cells]
        print(f'  Row{r}: {json.dumps(cells, ensure_ascii=False)}')
